import sqlite3
from datetime import datetime, date
from io import BytesIO
from decimal import Decimal, ROUND_HALF_UP

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from num2words import num2words

DB_PATH = "bmanager_facturation.db"


def get_conn():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_conn()
    cur = conn.cursor()

    # Clients
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS clients (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            email TEXT,
            phone TEXT
        )
        """
    )

    # Caisses / comptes
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS cash_accounts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL
        )
        """
    )

    # Config société (avec signataire)
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS company_config (
            id INTEGER PRIMARY KEY,
            name TEXT,
            legal_name TEXT,
            address TEXT,
            phone TEXT,
            email TEXT,
            currency TEXT DEFAULT '$',
            footer TEXT,
            logo BLOB,
            signatory_name TEXT,
            signatory_title TEXT
        )
        """
    )

    # Factures / proformas
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS invoices (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            client_id INTEGER NOT NULL,
            date TEXT NOT NULL,
            due_date TEXT,
            doc_type TEXT DEFAULT 'FACTURE',
            doc_number TEXT,
            footer TEXT,
            currency TEXT DEFAULT '$',
            tva_rate REAL DEFAULT 0,
            total_ht REAL NOT NULL,
            tva_amount REAL NOT NULL,
            total_ttc REAL NOT NULL,
            note TEXT,
            FOREIGN KEY(client_id) REFERENCES clients(id)
        )
        """
    )

    # Lignes de facture
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS invoice_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            invoice_id INTEGER NOT NULL,
            description TEXT NOT NULL,
            amount REAL NOT NULL,
            FOREIGN KEY(invoice_id) REFERENCES invoices(id)
        )
        """
    )

    # Paiements
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS payments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            invoice_id INTEGER NOT NULL,
            date TEXT NOT NULL,
            amount REAL NOT NULL,
            cash_account_id INTEGER,
            receiver TEXT,
            note TEXT,
            FOREIGN KEY(invoice_id) REFERENCES invoices(id),
            FOREIGN KEY(cash_account_id) REFERENCES cash_accounts(id)
        )
        """
    )

    # Migrations douces
    for alter in [
        "ALTER TABLE invoices ADD COLUMN doc_type TEXT DEFAULT 'FACTURE'",
        "ALTER TABLE invoices ADD COLUMN doc_number TEXT",
        "ALTER TABLE company_config ADD COLUMN signatory_name TEXT",
        "ALTER TABLE company_config ADD COLUMN signatory_title TEXT",
    ]:
        try:
            cur.execute(alter)
        except sqlite3.OperationalError:
            pass

    conn.commit()
    conn.close()


# ---------- CONFIG SOCIETE ----------

def get_company_config():
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT * FROM company_config WHERE id = 1")
    row = cur.fetchone()
    conn.close()
    return row


def save_company_config(
    name,
    legal_name,
    address,
    phone,
    email,
    currency,
    footer,
    logo_bytes,
    signatory_name,
    signatory_title,
):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT id FROM company_config WHERE id = 1")
    exists = cur.fetchone() is not None

    if exists:
        if logo_bytes is not None:
            cur.execute(
                """
                UPDATE company_config
                SET name=?, legal_name=?, address=?, phone=?,
                    email=?, currency=?, footer=?, logo=?,
                    signatory_name=?, signatory_title=?
                WHERE id=1
                """,
                (
                    name,
                    legal_name,
                    address,
                    phone,
                    email,
                    currency,
                    footer,
                    logo_bytes,
                    signatory_name,
                    signatory_title,
                ),
            )
        else:
            cur.execute(
                """
                UPDATE company_config
                SET name=?, legal_name=?, address=?, phone=?,
                    email=?, currency=?, footer=?,
                    signatory_name=?, signatory_title=?
                WHERE id=1
                """,
                (
                    name,
                    legal_name,
                    address,
                    phone,
                    email,
                    currency,
                    footer,
                    signatory_name,
                    signatory_title,
                ),
            )
    else:
        cur.execute(
            """
            INSERT INTO company_config
            (id, name, legal_name, address, phone, email,
             currency, footer, logo, signatory_name, signatory_title)
            VALUES (1, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                name,
                legal_name,
                address,
                phone,
                email,
                currency,
                footer,
                logo_bytes,
                signatory_name,
                signatory_title,
            ),
        )

    conn.commit()
    conn.close()


# ---------- CLIENTS / CAISSES ----------

def create_client_if_not_exists(name):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT id FROM clients WHERE name = ?", (name,))
    row = cur.fetchone()
    if row:
        client_id = row["id"]
    else:
        cur.execute("INSERT INTO clients (name) VALUES (?)", (name,))
        client_id = cur.lastrowid
        conn.commit()
    conn.close()
    return client_id


def get_clients():
    conn = get_conn()
    df = pd.read_sql_query("SELECT * FROM clients ORDER BY name", conn)
    conn.close()
    return df


def get_cash_accounts():
    conn = get_conn()
    df = pd.read_sql_query("SELECT * FROM cash_accounts ORDER BY name", conn)
    conn.close()
    return df


def create_cash_account_if_not_exists(name):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT id FROM cash_accounts WHERE name = ?", (name,))
    row = cur.fetchone()
    if row:
        acc_id = row["id"]
    else:
        cur.execute("INSERT INTO cash_accounts (name) VALUES (?)", (name,))
        acc_id = cur.lastrowid
        conn.commit()
    conn.close()
    return acc_id


# ---------- FACTURES / PROFORMAS ----------

def generate_invoice_number(doc_type: str, date_str: str) -> str:
    conn = get_conn()
    cur = conn.cursor()
    ymd = date_str.replace("-", "")
    prefix = "PRO-FORMA" if doc_type == "PROFORMA" else "FACT"
    cur.execute(
        "SELECT COUNT(*) AS c FROM invoices WHERE doc_type = ? AND date = ?",
        (doc_type, date_str),
    )
    row = cur.fetchone()
    count = row["c"] if row else 0
    conn.close()
    seq = count + 1
    return f"{prefix}-{ymd}-{seq:03d}"


def create_invoice(
    client_id,
    date_str,
    due_date_str,
    footer,
    currency,
    tva_rate,
    items,
    note="",
    doc_type="FACTURE",
):
    subtotal = sum(m for _, m in items)
    tva_amount = round(subtotal * (tva_rate / 100.0), 2) if tva_rate > 0 else 0.0
    total_ttc = round(subtotal + tva_amount, 2)
    doc_number = generate_invoice_number(doc_type, date_str)

    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO invoices (
            client_id, date, due_date,
            doc_type, doc_number,
            footer, currency,
            tva_rate, total_ht, tva_amount, total_ttc,
            note
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            client_id,
            date_str,
            due_date_str,
            doc_type,
            doc_number,
            footer,
            currency,
            tva_rate,
            subtotal,
            tva_amount,
            total_ttc,
            note,
        ),
    )
    invoice_id = cur.lastrowid

    for descr, amount in items:
        cur.execute(
            """
            INSERT INTO invoice_items (invoice_id, description, amount)
            VALUES (?, ?, ?)
            """,
            (invoice_id, descr, amount),
        )

    conn.commit()
    conn.close()
    return invoice_id


def add_payment(invoice_id, date_str, amount, cash_account_id, receiver, note=""):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO payments (invoice_id, date, amount, cash_account_id, receiver, note)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (invoice_id, date_str, amount, cash_account_id, receiver, note),
    )
    payment_id = cur.lastrowid
    conn.commit()
    conn.close()
    return payment_id


def compute_invoice_status(invoice_row):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        "SELECT COALESCE(SUM(amount), 0) AS paid FROM payments WHERE invoice_id = ?",
        (invoice_row["id"],),
    )
    paid = cur.fetchone()["paid"]
    conn.close()

    total = invoice_row["total_ttc"]
    balance = round(total - paid, 2)

    if paid <= 0:
        status = "Non payé"
    elif balance <= 0:
        status = "Payé"
    else:
        status = "Acompte (reste dû)"
    return total, paid, balance, status


def get_invoices(filters=None):
    conn = get_conn()
    base_query = "SELECT * FROM invoices"
    params = []
    where_clauses = []
    if filters:
        if filters.get("client_id"):
            where_clauses.append("client_id = ?")
            params.append(filters["client_id"])
        if filters.get("date_min"):
            where_clauses.append("date >= ?")
            params.append(filters["date_min"])
        if filters.get("date_max"):
            where_clauses.append("date <= ?")
            params.append(filters["date_max"])
    if where_clauses:
        base_query += " WHERE " + " AND ".join(where_clauses)
    base_query += " ORDER BY date DESC, id DESC"
    cur = conn.cursor()
    cur.execute(base_query, params)
    rows = cur.fetchall()
    conn.close()
    return rows


def get_invoice_with_items(invoice_id):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT * FROM invoices WHERE id = ?", (invoice_id,))
    inv = cur.fetchone()
    cur.execute(
        "SELECT * FROM invoice_items WHERE invoice_id = ? ORDER BY id", (invoice_id,)
    )
    items = cur.fetchall()
    conn.close()
    return inv, items


def update_invoice(
    invoice_id, date_str, due_date_str, footer, currency, tva_rate, items, note=""
):
    subtotal = sum(m for _, m in items)
    tva_amount = round(subtotal * (tva_rate / 100.0), 2) if tva_rate > 0 else 0.0
    total_ttc = round(subtotal + tva_amount, 2)

    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        UPDATE invoices
        SET date=?, due_date=?, footer=?, currency=?,
            tva_rate=?, total_ht=?, tva_amount=?, total_ttc=?, note=?
        WHERE id=?
        """,
        (
            date_str,
            due_date_str,
            footer,
            currency,
            tva_rate,
            subtotal,
            tva_amount,
            total_ttc,
            note,
            invoice_id,
        ),
    )

    cur.execute("DELETE FROM invoice_items WHERE invoice_id = ?", (invoice_id,))
    for descr, amount in items:
        cur.execute(
            """
            INSERT INTO invoice_items (invoice_id, description, amount)
            VALUES (?, ?, ?)
            """,
            (invoice_id, descr, amount),
        )

    conn.commit()
    conn.close()


def delete_invoice(invoice_id):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("DELETE FROM payments WHERE invoice_id = ?", (invoice_id,))
    cur.execute("DELETE FROM invoice_items WHERE invoice_id = ?", (invoice_id,))
    cur.execute("DELETE FROM invoices WHERE id = ?", (invoice_id,))
    conn.commit()
    conn.close()


def get_payments_for_invoice(invoice_id):
    conn = get_conn()
    df = pd.read_sql_query(
        "SELECT * FROM payments WHERE invoice_id=? ORDER BY date, id",
        conn,
        params=(invoice_id,),
    )
    conn.close()
    return df


def delete_payment(payment_id):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("DELETE FROM payments WHERE id=?", (payment_id,))
    conn.commit()
    conn.close()


# ---------- MONTANT EN LETTRES (corrigé) ----------

def amount_to_words_fr(amount: float) -> str:
    """
    101.16 -> 'cent un et seize centimes'
    50.00  -> 'cinquante'
    """
    value = Decimal(str(amount)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    entier = int(value)
    cents = int((value - Decimal(entier)) * 100)

    words_int = num2words(entier, lang="fr")
    if cents > 0:
        words_cents = num2words(cents, lang="fr")
        return f"{words_int} et {words_cents} centimes"
    return words_int


# ---------- DOCX HELPERS ----------

def _shade_cell(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def build_invoice_doc(invoice_id):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT * FROM invoices WHERE id=?", (invoice_id,))
    inv = cur.fetchone()
    cur.execute("SELECT * FROM clients WHERE id=?", (inv["client_id"],))
    client = cur.fetchone()
    cur.execute("SELECT * FROM invoice_items WHERE invoice_id=?", (invoice_id,))
    items = cur.fetchall()
    conn.close()

    cfg = get_company_config()
    total, paid, balance, status = compute_invoice_status(inv)

    try:
        date_display = datetime.strptime(inv["date"], "%Y-%m-%d").strftime("%d/%m/%Y")
    except Exception:
        date_display = inv["date"]

    if inv["doc_type"] == "PROFORMA":
        title_text = "FACTURE PRO FORMA"
        num_label = "N° proforma : "
    else:
        title_text = "FACTURE"
        num_label = "N° facture : "
    num_value = inv["doc_number"] or str(inv["id"])

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # En-tête logo + infos
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = header_table.rows[0].cells

    if cfg and cfg["logo"]:
        logo_stream = BytesIO(cfg["logo"])
        p_logo = left.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.add_run().add_picture(logo_stream, width=Inches(1.5))
    else:
        left.paragraphs[0].add_run("")
    _shade_cell(left, "FFFFFF")

    company_name = cfg["name"] if cfg and cfg["name"] else "b-manager"
    company_legal = (
        cfg["legal_name"]
        if cfg and cfg["legal_name"]
        else "Système de gestion & services IT"
    )
    company_address = cfg["address"] if cfg and cfg["address"] else ""
    company_phone = cfg["phone"] if cfg and cfg["phone"] else ""

    p1 = right.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run(company_name)
    r1.bold = True
    r1.font.size = Pt(12)

    p2 = right.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run(company_legal)

    if company_address:
        p3 = right.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.add_run(company_address)

    if company_phone:
        p4 = right.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p4.add_run(f"Tél : {company_phone}")

    p5 = right.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p5.add_run(f"Date : {date_display}")

    p6 = right.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p6.add_run(f"{num_label}{num_value}")
    _shade_cell(right, "FFFFFF")

    # Titre
    doc.add_paragraph("")
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title_para.add_run(title_text)
    rt.bold = True
    rt.font.size = Pt(16)
    rt.font.color.rgb = RGBColor(178, 34, 34)

    # Bloc infos
    doc.add_paragraph("")
    info_table = doc.add_table(rows=1, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_left, c_right = info_table.rows[0].cells
    _shade_cell(c_left, "F8FBFF")
    p_cli_title = c_left.paragraphs[0]
    p_cli_title.add_run("Facturer à :\n").bold = True
    p_cli_title.add_run(client["name"])

    _shade_cell(c_right, "F8FBFF")
    p_inf = c_right.paragraphs[0]
    p_inf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_inf.add_run("Informations facture\n").bold = True
    p_inf.add_run(f"Devise : {inv['currency']}\n")
    p_inf.add_run(f"TVA : {inv['tva_rate']:.2f} %\n")
    p_inf.add_run(f"Total TTC : {total:,.2f} {inv['currency']}")

    # Tableau lignes
    doc.add_paragraph("")
    items_table = doc.add_table(rows=1, cols=3)
    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = items_table.rows[0].cells
    hdr_cells[0].text = "N°"
    hdr_cells[1].text = "Désignation"
    hdr_cells[2].text = f"Montant ({inv['currency']})"
    for cell in hdr_cells:
        _shade_cell(cell, "B91C1C")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True

    subtotal = 0.0
    for i, row in enumerate(items, start=1):
        row_cells = items_table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = row["description"]
        row_cells[2].text = f"{row['amount']:,.2f}"
        subtotal += float(row["amount"])
        fill = "FFFFFF" if i % 2 == 1 else "FDECEC"
        for c in row_cells:
            _shade_cell(c, fill)

    # Totaux
    doc.add_paragraph("")
    totals_table = doc.add_table(rows=5, cols=2)
    totals_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    labels = [
        "Sous-total",
        f"TVA ({inv['tva_rate']:.2f} %)",
        "TOTAL TTC",
        "Montant payé",
        "Reste dû",
    ]
    values = [
        f"{subtotal:,.2f} {inv['currency']}",
        f"{inv['tva_amount']:,.2f} {inv['currency']}",
        f"{total:,.2f} {inv['currency']}",
        f"{paid:,.2f} {inv['currency']}",
        f"{balance:,.2f} {inv['currency']}",
    ]
    row_colors = ["FCA5A5", "FECACA", "B91C1C", "FFFFFF", "FFFFFF"]
    white_text = [False, False, True, False, False]

    for i in range(5):
        c_label, c_val = totals_table.rows[i].cells
        c_label.text = labels[i]
        c_val.text = values[i]
        _shade_cell(c_label, row_colors[i])
        _shade_cell(c_val, row_colors[i])
        for p in c_label.paragraphs + c_val.paragraphs:
            if white_text[i]:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(255, 255, 255)
            if i == 2:
                for r in p.runs:
                    r.bold = True
        c_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        c_val.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Montant en lettres (corrigé)
    doc.add_paragraph("")
    try:
        total_words = amount_to_words_fr(total)
        p_words = doc.add_paragraph()
        p_words.add_run("Montant en toutes lettres : ").bold = True
        p_words.add_run(f"{total_words} {inv['currency']}")
    except Exception:
        pass

    # Pied de page
    if inv["footer"]:
        doc.add_paragraph("")
        p_footer = doc.add_paragraph(inv["footer"])
        p_footer.style = "Intense Quote"

    # Signatures
    sign_name = cfg["signatory_name"] if cfg and cfg["signatory_name"] else ""
    sign_title = cfg["signatory_title"] if cfg and cfg["signatory_title"] else ""

    doc.add_paragraph("")
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_left, s_right = sig_table.rows[0].cells

    p_sl = s_left.paragraphs[0]
    p_sl.add_run("Signature du client").bold = True
    s_left.add_paragraph("\n\n\n")

    p_sr = s_right.paragraphs[0]
    p_sr.add_run("Pour la société").bold = True
    if sign_name:
        p_name = s_right.add_paragraph()
        p_name.add_run(sign_name).bold = True
    if sign_title:
        s_right.add_paragraph(sign_title)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def build_receipt_doc(payment_id):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT * FROM payments WHERE id=?", (payment_id,))
    pay = cur.fetchone()
    cur.execute("SELECT * FROM invoices WHERE id=?", (pay['invoice_id'],))
    inv = cur.fetchone()
    cur.execute("SELECT * FROM clients WHERE id=?", (inv['client_id'],))
    client = cur.fetchone()
    cur.execute("SELECT name FROM cash_accounts WHERE id=?", (pay['cash_account_id'],))
    cash_row = cur.fetchone()
    cash_name = cash_row["name"] if cash_row else "N/A"
    conn.close()

    cfg = get_company_config()
    total, paid_total, balance, status = compute_invoice_status(inv)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("REÇU DE PAIEMENT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(34, 139, 34)

    info = doc.add_paragraph()
    info.add_run("Date du paiement : ").bold = True
    info.add_run(f"{pay['date']}\n")
    info.add_run("Client : ").bold = True
    info.add_run(client["name"] + "\n")
    info.add_run("Document : ").bold = True
    info.add_run(f"{inv['doc_type']} {inv['doc_number'] or inv['id']}\n")
    info.add_run("Mode / Caisse : ").bold = True
    info.add_run(cash_name + "\n")
    if pay["receiver"]:
        info.add_run("Reçu par : ").bold = True
        info.add_run(pay["receiver"] + "\n")

    doc.add_paragraph("")
    doc.add_paragraph(f"Montant payé : {pay['amount']:,.2f} {inv['currency']}")
    doc.add_paragraph(
        f"Total facture : {total:,.2f} {inv['currency']} | "
        f"Total payé : {paid_total:,.2f} {inv['currency']} | "
        f"Reste dû : {balance:,.2f} {inv['currency']} (Statut : {status})"
    )

    sign_name = cfg["signatory_name"] if cfg and cfg["signatory_name"] else ""
    sign_title = cfg["signatory_title"] if cfg and cfg["signatory_title"] else ""

    if sign_name:
        doc.add_paragraph("")
        p_sig = doc.add_paragraph()
        p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p_sig.add_run("Pour la société : ")
        r.bold = True
        p_sig.add_run(sign_name)
        if sign_title:
            p_sig.add_run(f", {sign_title}")
        doc.add_paragraph("\n\n")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------- PAGES STREAMLIT ----------

def page_configuration():
    st.header("⚙️ Configuration de l'entreprise")
    cfg = get_company_config()

    name_default = cfg["name"] if cfg and cfg["name"] else "b-manager"
    legal_default = (
        cfg["legal_name"]
        if cfg and cfg["legal_name"]
        else "Système de gestion & services IT"
    )
    addr_default = cfg["address"] if cfg and cfg["address"] else ""
    phone_default = cfg["phone"] if cfg and cfg["phone"] else ""
    email_default = cfg["email"] if cfg and cfg["email"] else ""
    curr_default = cfg["currency"] if cfg and cfg["currency"] else "$"
    footer_default = (
        cfg["footer"] if cfg and cfg["footer"] else "Merci pour votre confiance."
    )
    sign_default = cfg["signatory_name"] if cfg and cfg["signatory_name"] else ""
    sign_title_default = (
        cfg["signatory_title"] if cfg and cfg["signatory_title"] else ""
    )

    st.markdown('<div class="section-card">', unsafe_allow_html=True)

    name = st.text_input("Nom commercial", name_default)
    legal_name = st.text_input("Raison sociale / Tagline", legal_default)
    address = st.text_area("Adresse", addr_default)
    col1, col2 = st.columns(2)
    with col1:
        phone = st.text_input("Téléphone", phone_default)
    with col2:
        email = st.text_input("Email", email_default)
    currency = st.text_input("Devise par défaut", curr_default)
    footer = st.text_area("Pied de page par défaut", footer_default)

    st.markdown("---")
    st.subheader("Signature sur les documents")
    col_sig1, col_sig2 = st.columns(2)
    with col_sig1:
        signatory_name = st.text_input(
            "Nom du signataire (ex: SERGE MOKELO)", sign_default
        )
    with col_sig2:
        signatory_title = st.text_input(
            "Fonction du signataire (ex: DG, Responsable IT)", sign_title_default
        )

    st.markdown("---")
    st.markdown("**Logo actuel :**")
    logo_bytes = None
    if cfg and cfg["logo"]:
        st.image(cfg["logo"], width=120)
    new_logo_file = st.file_uploader(
        "Changer le logo (PNG/JPG)", type=["png", "jpg", "jpeg"]
    )
    if new_logo_file is not None:
        logo_bytes = new_logo_file.read()

    if st.button("💾 Enregistrer la configuration"):
        save_company_config(
            name,
            legal_name,
            address,
            phone,
            email,
            currency,
            footer,
            logo_bytes,
            signatory_name,
            signatory_title,
        )
        st.success("Configuration enregistrée.")

    st.markdown("</div>", unsafe_allow_html=True)


def page_clients():
    st.header("👥 Clients")
    st.markdown('<div class="section-card">', unsafe_allow_html=True)

    df = get_clients()

    options = ["[Nouveau client]"] + df["name"].tolist() if not df.empty else [
        "[Nouveau client]"
    ]
    choice = st.selectbox("Sélectionner un client", options)

    if choice == "[Nouveau client]":
        name = st.text_input("Nom du client")
        email = st.text_input("Email")
        phone = st.text_input("Téléphone")
        if st.button("💾 Enregistrer le client"):
            if not name:
                st.error("Le nom est obligatoire.")
            else:
                cid = create_client_if_not_exists(name)
                conn = get_conn()
                cur = conn.cursor()
                cur.execute(
                    "UPDATE clients SET email=?, phone=? WHERE id=?",
                    (email, phone, cid),
                )
                conn.commit()
                conn.close()
                st.success("Client enregistré / mis à jour.")
    else:
        row = df[df["name"] == choice].iloc[0]
        name = st.text_input("Nom du client", row["name"])
        email = st.text_input("Email", row["email"] or "")
        phone = st.text_input("Téléphone", row["phone"] or "")
        colA, colB = st.columns(2)
        with colA:
            if st.button("💾 Mettre à jour ce client"):
                conn = get_conn()
                cur = conn.cursor()
                cur.execute(
                    "UPDATE clients SET name=?, email=?, phone=? WHERE id=?",
                    (name, email, phone, int(row["id"])),
                )
                conn.commit()
                conn.close()
                st.success("Client mis à jour.")
        with colB:
            if st.button("🗑️ Supprimer ce client"):
                conn = get_conn()
                cur = conn.cursor()
                cur.execute("DELETE FROM clients WHERE id=?", (int(row["id"]),))
                conn.commit()
                conn.close()
                st.success("Client supprimé.")

    st.markdown("---")
    st.subheader("Liste des clients")
    if df.empty:
        st.info("Aucun client pour l'instant.")
    else:
        st.dataframe(df)

    st.markdown("</div>", unsafe_allow_html=True)


def page_caisses():
    st.header("🏦 Caisses / comptes")
    st.markdown('<div class="section-card">', unsafe_allow_html=True)

    df = get_cash_accounts()
    options = ["[Nouvelle caisse / compte]"] + df["name"].tolist() if not df.empty else [
        "[Nouvelle caisse / compte]"
    ]
    choice = st.selectbox("Sélectionner une caisse / un compte", options)

    if choice == "[Nouvelle caisse / compte]":
        name = st.text_input("Nom de la caisse / du compte")
        if st.button("💾 Enregistrer la caisse / le compte"):
            if not name:
                st.error("Le nom est obligatoire.")
            else:
                create_cash_account_if_not_exists(name)
                st.success("Caisse / compte enregistré.")
    else:
        row = df[df["name"] == choice].iloc[0]
        name = st.text_input("Nom de la caisse / du compte", row["name"])
        colA, colB = st.columns(2)
        with colA:
            if st.button("💾 Mettre à jour"):
                conn = get_conn()
                cur = conn.cursor()
                cur.execute(
                    "UPDATE cash_accounts SET name=? WHERE id=?",
                    (name, int(row["id"])),
                )
                conn.commit()
                conn.close()
                st.success("Caisse / compte mis à jour.")
        with colB:
            if st.button("🗑️ Supprimer"):
                conn = get_conn()
                cur = conn.cursor()
                cur.execute(
                    "DELETE FROM cash_accounts WHERE id=?", (int(row["id"]),)
                )
                conn.commit()
                conn.close()
                st.success("Caisse / compte supprimé.")

    st.markdown("---")
    st.subheader("Liste des caisses / comptes")
    if df.empty:
        st.info("Aucune caisse / compte pour l'instant.")
    else:
        st.dataframe(df)

    st.markdown("</div>", unsafe_allow_html=True)


def page_nouvelle_facture():
    st.header("🧾 Créer une facture / proforma")
    cfg = get_company_config()
    company_name = cfg["name"] if cfg and cfg["name"] else "b-manager"
    company_legal = (
        cfg["legal_name"]
        if cfg and cfg["legal_name"]
        else "Système de gestion & services IT"
    )
    company_address = cfg["address"] if cfg and cfg["address"] else ""

    st.markdown('<div class="section-card">', unsafe_allow_html=True)

    doc_type_choice = st.radio(
        "Type de document", ["Facture", "Proforma"], horizontal=True
    )
    internal_doc_type = "FACTURE" if doc_type_choice == "Facture" else "PROFORMA"

    # Client
    clients_df = get_clients()
    client_names = ["[Nouveau client]"] + clients_df["name"].tolist()
    choix_client = st.selectbox("Client", client_names)

    if choix_client == "[Nouveau client]":
        new_name = st.text_input("Nom du nouveau client")
        new_email = st.text_input("Email du client", "")
        new_phone = st.text_input("Téléphone du client", "")
        if not new_name:
            st.info("Saisissez le nom du client pour continuer.")
            st.markdown("</div>", unsafe_allow_html=True)
            return
        client_id = create_client_if_not_exists(new_name)
        client_display_name = new_name
        conn = get_conn()
        cur = conn.cursor()
        cur.execute(
            "UPDATE clients SET email=?, phone=? WHERE id=?",
            (new_email, new_phone, client_id),
        )
        conn.commit()
        conn.close()
        st.success(f"Client créé : {new_name}")
    else:
        client_row = clients_df[clients_df["name"] == choix_client].iloc[0]
        client_id = int(client_row["id"])
        client_display_name = choix_client

    col1, col2 = st.columns(2)
    with col1:
        date_facture = st.date_input("Date de facture", date.today())
        due_date = st.date_input("Échéance (optionnel)", date.today())
    with col2:
        default_currency = cfg["currency"] if cfg and cfg["currency"] else "$"
        currency = st.text_input("Devise", default_currency)
        tva_rate = st.number_input(
            "TVA (%)", min_value=0.0, max_value=100.0, value=1.16, step=0.01
        )

    # Lignes de facture
    st.subheader("Lignes de facture")

    if "nb_items" not in st.session_state:
        st.session_state.nb_items = 1

    items = []
    for i in range(st.session_state.nb_items):
        c1, c2 = st.columns([3, 1])
        with c1:
            desc = st.text_input(f"Désignation {i+1}", key=f"desc_{i}")
        with c2:
            amt = st.number_input(
                f"Montant {i+1}", min_value=0.0, value=0.0, step=1.0, key=f"amt_{i}"
            )
        if desc and amt > 0:
            items.append((desc, amt))

    col_add, col_remove = st.columns(2)
    with col_add:
        if st.button("➕ Ajouter une ligne"):
            st.session_state.nb_items += 1
    with col_remove:
        if st.session_state.nb_items > 1:
            if st.button("➖ Supprimer la dernière ligne"):
                st.session_state.nb_items = max(1, st.session_state.nb_items - 1)

    subtotal = 0.0
    tva_amount = 0.0
    total_ttc = 0.0
    total_words = "—"

    if items:
        subtotal = sum(m for _, m in items)
        tva_amount = round(subtotal * (tva_rate / 100.0), 2) if tva_rate > 0 else 0.0
        total_ttc = round(subtotal + tva_amount, 2)
        try:
            total_words = amount_to_words_fr(total_ttc)
        except Exception:
            total_words = "—"
        st.info(
            f"**Total TTC provisoire :** {total_ttc:,.2f} {currency}\n\n"
            f"_Montant en toutes lettres :_ {total_words} {currency}"
        )

        # Prévisualisation facture
        date_str_for_number = date_facture.strftime("%Y-%m-%d")
        number_preview = generate_invoice_number(internal_doc_type, date_str_for_number)
        if internal_doc_type == "PROFORMA":
            preview_label = "N° proforma (prévisionnel)"
            title_text = "FACTURE PRO FORMA"
        else:
            preview_label = "N° facture (prévisionnel)"
            title_text = "FACTURE"

        rows_html = ""
        for i, (desc, amt) in enumerate(items, start=1):
            bg = "#ffffff" if i % 2 == 1 else "#f3f6fb"
            rows_html += f"""
            <tr style="background:{bg};">
              <td style="padding:6px 10px; text-align:center; font-size:12px;">{i}</td>
              <td style="padding:6px 10px; font-size:12px;">{desc}</td>
              <td style="padding:6px 10px; text-align:right; font-size:12px;">
                {amt:,.2f} {currency}
              </td>
            </tr>
            """

        preview_html = f"""
        <div class="preview-card">
          <div class="preview-header">
            <div>
              <div class="preview-company">{company_name}</div>
              <div class="preview-tagline">{company_legal}</div>
              <div class="preview-address">{company_address}</div>
            </div>
            <div class="preview-doc-meta">
              <div class="preview-doc-title">{title_text}</div>
              <div class="preview-doc-line">{preview_label} : {number_preview}</div>
              <div class="preview-doc-line">Date : {date_facture.strftime("%d/%m/%Y")}</div>
            </div>
          </div>

          <div class="preview-info-row">
            <div>
              <div class="preview-info-title">Facturer à</div>
              <div class="preview-info-value">{client_display_name}</div>
            </div>
            <div style="text-align:right;">
              <div class="preview-info-title">Infos facture</div>
              <div class="preview-info-value-small">
                Devise : {currency}<br/>
                TVA : {tva_rate:.2f} %<br/>
                Total TTC : {total_ttc:,.2f} {currency}
              </div>
            </div>
          </div>

          <table class="preview-table">
            <thead>
              <tr>
                <th style="border-top-left-radius:8px;">N°</th>
                <th>Désignation</th>
                <th style="border-top-right-radius:8px;">Montant ({currency})</th>
              </tr>
            </thead>
            <tbody>
              {rows_html}
            </tbody>
          </table>

          <div class="preview-totals-row">
            <table class="preview-totals-table">
              <tr>
                <td>Sous-total</td>
                <td>{subtotal:,.2f} {currency}</td>
              </tr>
              <tr>
                <td>TVA ({tva_rate:.2f} %)</td>
                <td>{tva_amount:,.2f} {currency}</td>
              </tr>
              <tr class="preview-total-ttc">
                <td>TOTAL TTC</td>
                <td>{total_ttc:,.2f} {currency}</td>
              </tr>
            </table>
          </div>

          <div class="preview-words">
            <span class="preview-words-label">Montant en toutes lettres :</span>
            <span>{total_words} {currency}</span>
          </div>
        </div>
        """
        st.markdown("### 👀 Aperçu de la facture")
        st.markdown(preview_html, unsafe_allow_html=True)

    # Pied de page & note
    default_footer = (
        cfg["footer"] if cfg and cfg["footer"] else "Merci pour votre confiance."
    )
    footer = st.text_area("Pied de page (conditions, remerciements…)", default_footer)
    note = st.text_area("Note interne (optionnel)", "")

    # Paiement initial
    st.subheader("💳 Paiement initial (optionnel)")
    mode_paiement = st.selectbox(
        "Type de paiement initial",
        ["Aucun paiement maintenant", "Paiement total", "Acompte"],
    )

    pay_date = None
    amount_init = 0.0
    cash_account_id = None
    receiver = ""
    note_pay = ""
    mode_text = ""
    accounts_df = None
    choix_acc = None

    if mode_paiement != "Aucun paiement maintenant":
        colA, colB = st.columns(2)
        with colA:
            pay_date = st.date_input(
                "Date du paiement initial", date.today(), key="pay_init_date"
            )
            montant_defaut = float(total_ttc) if total_ttc > 0 else 0.0
            label_amt = (
                "Montant payé (par défaut = Total TTC)"
                if mode_paiement == "Paiement total"
                else "Montant de l'acompte"
            )
            amount_init = st.number_input(
                label_amt,
                min_value=0.0,
                value=montant_defaut,
                step=1.0,
                key="pay_init_amount",
            )
        with colB:
            accounts_df = get_cash_accounts()
            acc_names = ["[Nouvelle caisse / compte]"] + accounts_df["name"].tolist()
            choix_acc = st.selectbox(
                "Caisse / compte", acc_names, key="pay_init_acc"
            )
            receiver = st.text_input(
                "Reçu par (nom de la personne)", "", key="pay_init_receiver"
            )

        if choix_acc == "[Nouvelle caisse / compte]":
            new_acc = st.text_input(
                "Nom de la nouvelle caisse / compte",
                "Caisse principale",
                key="pay_init_newacc",
            )
            if new_acc:
                cash_account_id = create_cash_account_if_not_exists(new_acc)
        else:
            if accounts_df is not None and not accounts_df.empty:
                row_acc = accounts_df[accounts_df["name"] == choix_acc].iloc[0]
                cash_account_id = int(row_acc["id"])

        mode_text = st.text_input(
            "Mode de paiement (ex: Cash, M-Pesa, Banque)",
            "",
            key="pay_init_mode",
        )
        note_pay = st.text_area(
            "Note sur ce paiement initial (optionnel)",
            "",
            key="pay_init_note",
        )

    if st.button("💾 Enregistrer le document"):
        if not items:
            st.error("Ajoutez au moins une ligne.")
            st.markdown("</div>", unsafe_allow_html=True)
            return

        date_str = date_facture.strftime("%Y-%m-%d")
        due_str = due_date.strftime("%Y-%m-%d")

        invoice_id = create_invoice(
            client_id=client_id,
            date_str=date_str,
            due_date_str=due_str,
            footer=footer,
            currency=currency,
            tva_rate=tva_rate,
            items=items,
            note=note,
            doc_type=internal_doc_type,
        )
        st.success(f"Document enregistré avec l'ID {invoice_id}.")

        # Export DOCX
        doc_bytes = build_invoice_doc(invoice_id)
        st.download_button(
            "⬇️ Télécharger la facture / proforma (DOCX)",
            data=doc_bytes,
            file_name=f"document_{invoice_id}.docx",
            mime=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )

        # Paiement initial
        if (
            mode_paiement != "Aucun paiement maintenant"
            and amount_init > 0
            and cash_account_id
        ):
            full_note = note_pay
            if mode_text:
                full_note = (
                    f"Mode: {mode_text} - {full_note}"
                    if full_note
                    else f"Mode: {mode_text}"
                )
            payment_id = add_payment(
                invoice_id=invoice_id,
                date_str=pay_date.strftime("%Y-%m-%d"),
                amount=amount_init,
                cash_account_id=cash_account_id,
                receiver=receiver,
                note=full_note,
            )
            st.success(
                f"Paiement initial enregistré (ID paiement {payment_id}). "
                "Le statut du document est mis à jour automatiquement."
            )
            recu_bytes = build_receipt_doc(payment_id)
            st.download_button(
                "⬇️ Télécharger le reçu du paiement initial (DOCX)",
                data=recu_bytes,
                file_name=f"recu_{payment_id}.docx",
                mime=(
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
            )

    st.markdown("</div>", unsafe_allow_html=True)


def page_paiements():
    st.header("💳 Paiements & reçus")
    st.markdown('<div class="section-card">', unsafe_allow_html=True)

    rows = get_invoices(None)
    if not rows:
        st.info("Aucune facture enregistrée pour l'instant.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    labels = []
    for r in rows:
        total, paid, balance, status = compute_invoice_status(r)
        labels.append(
            f"{r['doc_type']} | {r['doc_number'] or r['id']} | "
            f"{r['date']} | Total {total:.2f} | {status}"
        )

    idx = st.selectbox(
        "Sélectionnez un document",
        range(len(rows)),
        format_func=lambda i: labels[i],
    )
    inv = rows[idx]
    total, paid, balance, status = compute_invoice_status(inv)

    st.write(f"**Document :** {inv['doc_type']} {inv['doc_number'] or inv['id']}")
    st.write(f"**Client ID :** {inv['client_id']}")
    st.write(f"**Total TTC :** {total:.2f}")
    st.write(f"**Total payé :** {paid:.2f}")
    st.write(f"**Reste dû :** {balance:.2f}")
    st.write(f"**Statut :** {status}")

    st.markdown("---")
    st.subheader("Enregistrer un paiement")

    col1, col2 = st.columns(2)
    with col1:
        pay_date = st.date_input("Date du paiement", date.today())
        default_amount = float(balance) if balance > 0 else 0.0
        amount = st.number_input(
            "Montant payé", min_value=0.0, value=default_amount, step=1.0
        )
    with col2:
        accounts_df = get_cash_accounts()
        acc_names = ["[Nouvelle caisse / compte]"] + accounts_df["name"].tolist()
        choix_acc = st.selectbox("Caisse / compte", acc_names)
        receiver = st.text_input("Reçu par (nom de la personne)", "")

    if choix_acc == "[Nouvelle caisse / compte]":
        new_acc = st.text_input(
            "Nom de la nouvelle caisse / compte", "Caisse principale"
        )
        if new_acc:
            cash_account_id = create_cash_account_if_not_exists(new_acc)
        else:
            cash_account_id = None
    else:
        row_acc = accounts_df[accounts_df["name"] == choix_acc].iloc[0]
        cash_account_id = int(row_acc["id"])

    mode_text = st.text_input("Mode de paiement (ex: Cash, M-Pesa, Banque)", "")
    note = st.text_area("Note sur le paiement (optionnel)", "")

    if st.button("💾 Enregistrer le paiement & générer le reçu"):
        if amount <= 0:
            st.error("Le montant doit être supérieur à 0.")
        else:
            full_note = note
            if mode_text:
                full_note = (
                    f"Mode: {mode_text} - {note}" if note else f"Mode: {mode_text}"
                )
            payment_id = add_payment(
                invoice_id=inv["id"],
                date_str=pay_date.strftime("%Y-%m-%d"),
                amount=amount,
                cash_account_id=cash_account_id,
                receiver=receiver,
                note=full_note,
            )
            st.success(f"Paiement enregistré (ID paiement {payment_id}).")
            receipt_bytes = build_receipt_doc(payment_id)
            st.download_button(
                "⬇️ Télécharger le reçu (DOCX)",
                data=receipt_bytes,
                file_name=f"recu_{payment_id}.docx",
                mime=(
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
            )

    st.markdown("---")
    st.subheader("🧾 Paiements enregistrés pour ce document")

    pay_df = get_payments_for_invoice(inv["id"])
    if pay_df.empty:
        st.info("Aucun paiement enregistré pour ce document.")
    else:
        st.dataframe(
            pay_df[
                ["id", "date", "amount", "cash_account_id", "receiver", "note"]
            ].rename(
                columns={
                    "id": "ID",
                    "date": "Date",
                    "amount": "Montant",
                    "cash_account_id": "Caisse / compte",
                    "receiver": "Reçu par",
                    "note": "Note",
                }
            )
        )
        options = ["Aucun"] + [
            f"{row['id']} - {row['date']} - {row['amount']}"
            for _, row in pay_df.iterrows()
        ]
        choice = st.selectbox("Sélectionner un paiement à supprimer", options)
        if choice != "Aucun":
            pay_id = int(choice.split(" - ")[0])
            if st.button("🗑️ Supprimer ce paiement"):
                delete_payment(pay_id)
                st.success(
                    "Paiement supprimé. Rafraîchissez la page pour mettre à jour les totaux."
                )

    st.markdown("</div>", unsafe_allow_html=True)


def page_factures():
    st.header("📂 Documents & filtres")
    st.markdown('<div class="section-card">', unsafe_allow_html=True)

    clients_df = get_clients()
    client_options = ["[Tous]"] + clients_df["name"].tolist()
    col1, col2, col3 = st.columns(3)
    with col1:
        client_choice = st.selectbox("Client", client_options)
    with col2:
        date_min = st.date_input("Date min", value=date(2024, 1, 1))
    with col3:
        date_max = st.date_input("Date max", value=date.today())

    filters = {
        "client_id": None,
        "date_min": date_min.strftime("%Y-%m-%d"),
        "date_max": date_max.strftime("%Y-%m-%d"),
    }
    if client_choice != "[Tous]":
        client_row = clients_df[clients_df["name"] == client_choice].iloc[0]
        filters["client_id"] = int(client_row["id"])

    rows = get_invoices(filters)
    if not rows:
        st.info("Aucun document ne correspond aux filtres.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    records = []
    for r in rows:
        total, paid, balance, status = compute_invoice_status(r)
        records.append(
            {
                "ID": r["id"],
                "Type": r["doc_type"],
                "N°": r["doc_number"],
                "Client ID": r["client_id"],
                "Date": r["date"],
                "Échéance": r["due_date"],
                "Total TTC": total,
                "Total payé": paid,
                "Reste dû": balance,
                "Statut": status,
            }
        )
    df = pd.DataFrame.from_records(records)

    etats = ["Tous", "Payé", "Non payé", "Acompte (reste dû)"]
    etat_choice = st.radio("Filtrer par état", etats, horizontal=True)
    if etat_choice != "Tous":
        df = df[df["Statut"] == etat_choice]

    st.dataframe(
        df.style.format(
            {
                "Total TTC": "{:,.2f}",
                "Total payé": "{:,.2f}",
                "Reste dû": "{:,.2f}",
            }
        )
    )

    st.markdown("**Documents par client**")
    group_client = df.groupby("Client ID")[["Total TTC", "Total payé", "Reste dû"]].sum()
    st.dataframe(group_client)

    st.markdown("---")
    st.subheader("✏️ Modifier ou supprimer un document")

    ids = df["ID"].tolist()
    if ids:
        selected_id = st.selectbox("Choisir un document à éditer", ids)

        if selected_id:
            inv, items = get_invoice_with_items(selected_id)
            with st.form(f"edit_invoice_{selected_id}"):
                col1, col2 = st.columns(2)
                with col1:
                    d_val = inv["date"] or date.today().isoformat()
                    date_facture = st.date_input(
                        "Date de facture", value=date.fromisoformat(d_val)
                    )
                    if inv["due_date"]:
                        due_val = date.fromisoformat(inv["due_date"])
                    else:
                        due_val = date_facture
                    due_date = st.date_input("Échéance", value=due_val)
                with col2:
                    currency = st.text_input("Devise", inv["currency"] or "$")
                    tva_rate = st.number_input(
                        "TVA (%)",
                        min_value=0.0,
                        max_value=100.0,
                        value=float(inv["tva_rate"] or 0),
                        step=0.01,
                    )
                footer = st.text_area("Pied de page", inv["footer"] or "")
                note = st.text_area("Note interne", inv["note"] or "")

                st.write("Lignes de facture")
                edited_items = []
                for row in items:
                    desc = st.text_input(
                        f"Désignation #{row['id']}",
                        value=row["description"],
                        key=f"edit_desc_{row['id']}",
                    )
                    amt = st.number_input(
                        f"Montant #{row['id']}",
                        min_value=0.0,
                        value=float(row["amount"]),
                        step=1.0,
                        key=f"edit_amt_{row['id']}",
                    )
                    if desc and amt > 0:
                        edited_items.append((desc, amt))

                colA, colB = st.columns(2)
                with colA:
                    do_update = st.form_submit_button("💾 Mettre à jour")
                with colB:
                    do_delete = st.form_submit_button("🗑️ Supprimer le document")

            if do_update:
                if not edited_items:
                    st.error("Ajoutez au moins une ligne avant d'enregistrer.")
                else:
                    update_invoice(
                        selected_id,
                        date_facture.strftime("%Y-%m-%d"),
                        due_date.strftime("%Y-%m-%d"),
                        footer,
                        currency,
                        tva_rate,
                        edited_items,
                        note,
                    )
                    st.success(
                        "Document mis à jour. Rafraîchissez la page pour voir les nouvelles valeurs."
                    )

            if do_delete:
                delete_invoice(selected_id)
                st.success(
                    "Document supprimé (lignes + paiements associés). Rafraîchissez la page."
                )

    st.markdown("</div>", unsafe_allow_html=True)


def page_rapports():
    st.header("📊 Rapports & caisse")
    st.markdown('<div class="section-card">', unsafe_allow_html=True)

    conn = get_conn()
    inv_df = pd.read_sql_query("SELECT * FROM invoices", conn)
    pay_df = pd.read_sql_query("SELECT * FROM payments", conn)
    acc_df = pd.read_sql_query("SELECT * FROM cash_accounts", conn)
    conn.close()

    if inv_df.empty:
        st.info("Aucune donnée pour le moment.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    if not pay_df.empty:
        pay_group = pay_df.groupby("invoice_id")["amount"].sum()
    else:
        pay_group = pd.Series(dtype=float)

    inv_df["total_paid"] = inv_df["id"].map(pay_group).fillna(0.0)
    inv_df["balance"] = inv_df["total_ttc"] - inv_df["total_paid"]

    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total facturé (TTC)", f"{inv_df['total_ttc'].sum():,.2f}")
    with col2:
        st.metric("Total payé", f"{inv_df['total_paid'].sum():,.2f}")
    with col3:
        st.metric("Reste dû global", f"{inv_df['balance'].sum():,.2f}")

    st.markdown("---")
    st.subheader("État de la caisse (par compte)")

    if pay_df.empty or acc_df.empty:
        st.info("Pas encore de paiements enregistrés.")
    else:
        pay_acc = pay_df.groupby("cash_account_id")["amount"].sum().reset_index()
        merged = pay_acc.merge(
            acc_df, left_on="cash_account_id", right_on="id", how="left"
        )
        merged = merged.rename(
            columns={"name": "Caisse / compte", "amount": "Total encaissé"}
        )
        st.dataframe(merged[["Caisse / compte", "Total encaissé"]])

    st.markdown("---")
    st.subheader("Graphique : paiements par mois")

    if not pay_df.empty:
        pay_df["date"] = pd.to_datetime(pay_df["date"])
        pay_df["mois"] = pay_df["date"].dt.to_period("M").astype(str)
        pay_month = pay_df.groupby("mois")["amount"].sum().reset_index()
        pay_month = pay_month.sort_values("mois")
        st.bar_chart(pay_month.set_index("mois"))
    else:
        st.info("Aucun paiement pour générer un graphique.")

    st.markdown("</div>", unsafe_allow_html=True)


# ---------- MAIN ----------

def main():
    st.set_page_config(
        page_title="b-manager - Facturation",
        page_icon="💼",
        layout="wide",
    )

    # Global style
    st.markdown(
        """
        <style>
        .main {
            background-color: #f5f7fb;
        }
        .stButton>button {
            border-radius: 999px;
            padding: 0.4rem 1.3rem;
            font-weight: 600;
        }
        .stTextInput>div>div>input, .stTextArea textarea {
            border-radius: 8px;
        }
        [data-testid="stSidebar"] {
            background: radial-gradient(circle at top left, #b91c1c 0, #111827 45%, #020617 100%);
            color: #f9fafb;
        }
        [data-testid="stSidebar"] * {
            color: #e5e7eb !important;
        }
        .sidebar-title {
            font-size: 0.9rem;
            font-weight: 600;
            letter-spacing: .08em;
            text-transform: uppercase;
            color: #9ca3af;
            margin-bottom: 0.4rem;
        }
        div[role="radiogroup"] > label {
            padding: 0.25rem 0.7rem;
            border-radius: 999px;
            margin-bottom: 0.2rem;
        }
        div[role="radiogroup"] > label[aria-checked="true"] {
            background: rgba(248, 250, 252, 0.18);
            box-shadow: 0 0 0 1px rgba(248, 250, 252, 0.35);
        }
        .section-card {
            background:#ffffff;
            border-radius:18px;
            padding:18px 22px;
            box-shadow:0 10px 26px rgba(15,23,42,0.06);
            margin-bottom:24px;
        }
        .preview-card {
          max-width:850px;
          margin:20px auto;
          padding:22px 26px;
          background:#ffffff;
          border-radius:14px;
          box-shadow:0 8px 22px rgba(15,23,42,0.08);
          font-family:Segoe UI, system-ui, -apple-system, BlinkMacSystemFont, sans-serif;
          color:#111827;
        }
        .preview-header {
          display:flex;
          justify-content:space-between;
          align-items:flex-start;
          margin-bottom:18px;
        }
        .preview-company {
          font-size:15px;
          font-weight:700;
          color:#111827;
          margin-bottom:2px;
        }
        .preview-tagline {
          font-size:11px;
          color:#6b7280;
        }
        .preview-address {
          font-size:11px;
          color:#6b7280;
          margin-top:4px;
        }
        .preview-doc-meta {
          text-align:right;
        }
        .preview-doc-title {
          font-size:22px;
          font-weight:800;
          letter-spacing:1px;
          color:#b91c1c;
        }
        .preview-doc-line {
          font-size:11px;
          color:#6b7280;
        }
        .preview-info-row {
          display:flex;
          justify-content:space-between;
          background:#f3f6fb;
          border-radius:10px;
          padding:10px 14px;
          margin-bottom:14px;
        }
        .preview-info-title {
          font-size:11px;
          font-weight:700;
          color:#374151;
          margin-bottom:4px;
        }
        .preview-info-value {
          font-size:12px;
          color:#111827;
        }
        .preview-info-value-small {
          font-size:11px;
          color:#4b5563;
        }
        .preview-table {
          width:100%;
          border-collapse:collapse;
          margin-top:4px;
          margin-bottom:10px;
        }
        .preview-table th {
          background:#b91c1c;
          color:#ffffff;
          padding:6px 10px;
          font-size:11px;
          text-align:left;
        }
        .preview-table td {
          padding:6px 10px;
          font-size:12px;
        }
        .preview-table tbody tr:nth-child(even) {
          background:#f3f6fb;
        }
        .preview-totals-row {
          display:flex;
          justify-content:flex-end;
          margin-top:4px;
        }
        .preview-totals-table {
          font-size:11px;
          border-collapse:collapse;
          min-width:260px;
        }
        .preview-totals-table td {
          padding:4px 10px;
          text-align:right;
          background:#ffffff;
        }
        .preview-totals-table .preview-total-ttc td {
          background:#b91c1c;
          color:#ffffff;
          font-weight:700;
        }
        .preview-words {
          margin-top:12px;
          padding:8px 12px;
          border-radius:8px;
          background:#f9fafb;
          font-size:11px;
          color:#374151;
        }
        .preview-words-label {
          font-weight:600;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.title("💼 b-manager – Mini application de facturation")
    init_db()

    st.sidebar.markdown(
        '<div class="sidebar-title">Navigation</div>', unsafe_allow_html=True
    )
    menu = st.sidebar.radio(
        "",
        [
            "👥 Clients",
            "🏦 Caisses / comptes",
            "💳 Paiements & reçus",
            "🧾 Créer une facture / proforma",
            "📂 Documents & filtres",
            "📊 Rapports & caisse",
            "⚙️ Configuration",
        ],
    )

    if "Clients" in menu:
        page_clients()
    elif "Caisses" in menu:
        page_caisses()
    elif "Créer une facture" in menu:
        page_nouvelle_facture()
    elif "Paiements" in menu:
        page_paiements()
    elif "Documents" in menu:
        page_factures()
    elif "Rapports" in menu:
        page_rapports()
    elif "Configuration" in menu:
        page_configuration()


if __name__ == "__main__":
    main()
